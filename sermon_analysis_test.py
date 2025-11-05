
import os
from docx import Document
self_path = os.path.abspath(__file__)
base_path = os.path.dirname(self_path)
wordfile_path = os.path.join(base_path, "20251019新竹主日週報.docx")
wordfile_path = os.path.join(base_path, "20250928新竹主日週報.docx")
# wordfile_path = os.path.join(base_path, "202501005新竹主日週報.docx")
# wordfile_path = os.path.join(base_path, "20251012新竹主日週報.docx")
doc = Document(wordfile_path)
sermon = []
abbr_to_full = {
    "創": "創世記",
    "出": "出埃及記",
    "利": "利未記",
    "民": "民數記",
    "申": "申命記",
    "書": "約書亞記",
    "士": "士師記",
    "得": "路得記",
    "撒上": "撒母耳記上",
    "撒下": "撒母耳記下",
    "王上": "列王紀上",
    "王下": "列王紀下",
    "代上": "歷代志上",
    "代下": "歷代志下",
    "拉": "以斯拉記",
    "尼": "尼希米記",
    "斯": "以斯帖記",
    "伯": "約伯記",
    "詩": "詩篇",
    "箴": "箴言",
    "傳": "傳道書",
    "歌": "雅歌",
    "賽": "以賽亞書",
    "耶": "耶利米書",
    "哀": "耶利米哀歌",
    "結": "以西結書",
    "但": "但以理書",
    "何": "何西阿書",
    "珥": "約珥書",
    "摩": "阿摩司書",
    "俄": "俄巴底亞書",
    "拿": "約拿書",
    "彌": "彌迦書",
    "鴻": "何西阿書",  # 小先知書，部分版本略有不同
    "哈": "哈巴谷書",
    "番": "西番雅書",
    "該": "哈該書",
    "瑪": "撒迦利亞書",
    "亞": "瑪拉基書",
    "太": "馬太福音",
    "可": "馬可福音",
    "路": "路加福音",
    "約": "約翰福音",
    "徒": "使徒行傳",
    "羅": "羅馬書",
    "林前": "哥林多前書",
    "林後": "哥林多後書",
    "加": "加拉太書",
    "弗": "以弗所書",
    "腓": "腓立比書",
    "西": "歌羅西書",
    "帖前": "帖撒羅尼迦前書",
    "帖後": "帖撒羅尼迦後書",
    "提前": "提摩太前書",
    "提後": "提摩太後書",
    "多": "提多書",
    "門": "腓利門書",
    "來": "希伯來書",
    "雅": "雅各書",
    "彼前": "彼得前書",
    "彼後": "彼得後書",
    "約壹": "約翰一書",
    "約貳": "約翰二書",
    "約參": "約翰三書",
    "猶": "猶大書",
    "啟": "啟示錄"
}

# 逐個表格抓文字
for t_idx, table in enumerate(doc.tables):
    # print(f"=== 表格 {t_idx+1} ===")
    for r_idx, row in enumerate(table.rows):
        # 取每個儲存格文字，去掉前後空白
        
        # print(row_texts)
        tatil = row.cells[0].text.strip()
        if tatil == "證道": #添加全形分割符號支援
            bold_texts  = ""
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if run.bold and text:
                            print("加粗文字：", bold_texts)
                            if bold_texts == "證道":
                                bold_texts = text
                            else:
                                if text in ["錢致榮", "牧師", "傳道", "吳佩倫"]:
                                    if bold_texts != "":
                                        # print("遇到講員名稱，結束證道內容擷取", bold_texts)
                                        sermon.append(bold_texts)
                                    bold_texts = ""
                                    continue
                                for symbol in ["、", ".", ")"]:
                                    if symbol in bold_texts:
                                        print("遇到分隔符號，結束證道內容擷取", symbol, bold_texts)
                                        first_part = bold_texts.split(symbol)[0]
                                        second_part = first_part[-1][-1] + symbol + bold_texts.split(symbol)[1]
                                        first_part = first_part[:-1]
                                        if first_part:
                                            sermon.append(first_part)
                                        if second_part: 
                                            sermon.append(second_part)
                                        bold_texts = text
                                        break
                                else:
                                    for book in ["創", "出", "利", "民", "申", "書", "士", "得", "撒上", "撒下", "王上", "王下", "代上", "代下", "拉", "尼", "斯", "伯", "詩", "箴", "傳", "歌", "賽", "耶", "哀", "結", "但", "何", "珥", "摩", "俄", "拿", "彌", "鴻", "哈", "番", "該", "瑪", "亞", "太", "可", "路", "約", "徒", "羅", "林前", "林後", "加", "弗", "腓", "西", "帖前", "帖後", "提前", "提後", "多", "門", "來", "雅", "彼前", "彼後", "約壹", "約貳", "約參", "猶", "啟"]:
                                        if book in bold_texts:
                                            print(bold_texts.index(book)+len(book), len(bold_texts))
                                        if book in bold_texts and (bold_texts.index(book)+len(book) == len(bold_texts) or bold_texts[bold_texts.index(book)+1] in ["1", "2", "3", "4", "5", "6", "7", "8", "9", "0"]):
                                            first_part = bold_texts.split(book)[0]
                                            bold_texts = bold_texts.replace(first_part, "")
                                            print("遇到書卷，結束證道內容擷取", book, first_part, bold_texts)
                                            if first_part:
                                                sermon.append(first_part)
                                            if bold_texts:
                                                sermon.append(bold_texts)
                                            bold_texts = text
                                            break
                                    else:
                                        n = ""
                                        min_num_index = len(bold_texts)
                                        for num in ["1", "2", "3", "4", "5", "6", "7", "8", "9", "0"]:
                                            if num in bold_texts:
                                                idx = bold_texts.index(num)
                                                if idx < min_num_index:
                                                    min_num_index = idx
                                                    n = num
                                        print("最小數字索引", min_num_index, n)
                                        if n and n in bold_texts:
                                            print(bold_texts.split(n))
                                            first_part = bold_texts.split(n)[0]
                                            bold_texts = bold_texts.replace(first_part, "")
                                            print("遇到數字，結束證道內容擷取", n, first_part)
                                            if first_part:
                                                sermon.append(first_part)
                                            bold_texts += text

                                        else:
                                            bold_texts += text
                                            continue
sermon.append(bold_texts)
print(sermon)


